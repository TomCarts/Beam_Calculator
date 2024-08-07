import streamlit as st
import pandas as pd
from PIL import Image
import plotly.express as px
from PyNite import FEModel3D
import io 
from docx.shared import Inches
import numpy as np
from docx import Document
import plotly.graph_objects as go


# Analysis function
def beam_analysis(nodes_df, members_df, point_loads_df):
# Create a new finite element model
    beam = FEModel3D()
    
    #Add material
    for i in range(len(material_df)):
        beam.add_material(material_df.iloc[i,0],material_df.iloc[i,1],material_df.iloc[i,2],material_df.iloc[i,3],material_df.iloc[i,4])
    
    # Add nodes 
    for i in range(len(nodes_df)):
        beam.add_node(nodes_df.iloc[i,0],nodes_df.iloc[i,1],nodes_df.iloc[i,2],nodes_df.iloc[i,3])

    # Add  beams:
    for i in range(len(members_df)):
        beam.add_member(members_df.iloc[i,0], members_df.iloc[i,1], members_df.iloc[i,2], members_df.iloc[i,3], members_df.iloc[i,4], members_df.iloc[i,5], members_df.iloc[i,6], members_df.iloc[i,7])

    # Provide simple supports
    for i in range(len(nodes_df)):
        beam.def_support(nodes_df.iloc[i,0], nodes_df.iloc[i,4], nodes_df.iloc[i,5], nodes_df.iloc[i,6], nodes_df.iloc[i,7], nodes_df.iloc[i,8], nodes_df.iloc[i,9])  

    # Add point loads to member
    for i in range(len(point_loads_df)):
        beam.add_member_pt_load(point_loads_df.iloc[i,0],point_loads_df.iloc[i,1],point_loads_df.iloc[i,2],point_loads_df.iloc[i,3]) # 5 kips Dead load

    #add_member_dist_load(member_name, Direction, w1, w2, x1=None, x2=None, case='Case 1')
    for i in range(len(udl_loads_df)):
        beam.add_member_dist_load(udl_loads_df.iloc[i,0],udl_loads_df.iloc[i,1],udl_loads_df.iloc[i,2],udl_loads_df.iloc[i,3],udl_loads_df.iloc[i,4],udl_loads_df.iloc[i,5]) # 5 kips Dead load4

    # Analyze the beam and perform a statics check
    beam.analyze(check_statics=True)
           
    return beam

# Display function for plots
def display_plots(beam,members_df):
    for i in range(len(members_df)):
        st.pyplot(beam.Members[members_df.iloc[i,0]].plot_shear('Fy'))
        st.pyplot(beam.Members[members_df.iloc[i,0]].plot_moment('Mz'))
        st.pyplot(beam.Members[members_df.iloc[i,0]].plot_deflection('dy'))
        
def results_table(beam,members_df,nodes_df):
    results_df=pd.DataFrame()
    Mz=[]
    Fy=[]
    dy=[]
    X=[]
    for i in range(len(members_df)):
        #for j in range((nodes_df.iloc[i+1,1]-nodes_df.iloc[i,1])*10):
        x=np.linspace(0,(nodes_df.iloc[i+1,1]-nodes_df.iloc[i,1]),100)
        for xs in x:
            Mz.append(beam.Members[members_df.iloc[i,0]].moment('Mz',xs))
            Fy.append(beam.Members[members_df.iloc[i,0]].shear('Fy',xs))
            dy.append(beam.Members[members_df.iloc[i,0]].deflection('dy',xs))
            X.append(xs+nodes_df.iloc[i,1])

    results_df['X']=X
    results_df['Mz']=Mz
    results_df['Fy']=Fy
    results_df['dy']=dy
    results_df['beam']=0  

    return results_df

# Free Body Diagram Visualization function
def plot_free_body_diagram(nodes_df, members_df, point_loads_df, udl_loads_df):
    fig = go.Figure()

    # Draw beams
    for index, row in members_df.iterrows():
        node_lhs = nodes_df[nodes_df['name'] == row['Node_LHS']]
        node_rhs = nodes_df[nodes_df['name'] == row['Node_RHS']]
        fig.add_trace(go.Scatter(x=[node_lhs['x'].values[0], node_rhs['x'].values[0]],
                                 y=[0, 0],
                                 mode='lines',
                                 line=dict(width=5),
                                 name=row['Name']))

    # Draw supports
    for index, row in nodes_df.iterrows():
        support_label = ''
        if row['R_Fx']:
            support_label += 'Fx, '
        if row['R_Fy']:
            support_label += 'Fy, '
        if row['R_Fz']:
            support_label += 'Fz, '
        if row['R_Mx']:
            support_label += 'Mx, '
        if row['R_My']:
            support_label += 'My, '
        if row['R_Mz']:
            support_label += 'Mz, '

        if support_label:
            fig.add_annotation(x=row['x'], y=-0.1, text=support_label[:-2], showarrow=False)

    # Draw point loads as arrows
    for index, load in point_loads_df.iterrows():
        member = load['Member']
        member_data = members_df[members_df['Name'] == member]
        load_x = load['x1'] + nodes_df[nodes_df['name'] == member_data['Node_LHS'].values[0]]['x'].values[0]
        load_value = load['Value']

        fig.add_annotation(x=load_x, y=0, text=f"{load_value} kN",
                           showarrow=True,
                           arrowhead=2,
                           arrowsize=1,
                           arrowwidth=2,
                           arrowcolor="blue",
                           ax=0,
                           ay=-30)

    # Draw UDLs as single thick line with label
    for index, load in udl_loads_df.iterrows():
        member = load['Member']
        member_data = members_df[members_df['Name'] == member]
        load_x1 = load['x1'] + nodes_df[nodes_df['name'] == member_data['Node_LHS'].values[0]]['x'].values[0]
        load_x2 = load['x2'] + nodes_df[nodes_df['name'] == member_data['Node_LHS'].values[0]]['x'].values[0]
        load_value = load['w1']  # Assuming w1 represents the value of UDL
        udl_center = (load_x1 + load_x2) / 2
        
        fig.add_shape(type="line",
                      x0=load_x1,
                      y0=0.3,  # Adjust the height of the UDL line above the beam as needed
                      x1=load_x2,
                      y1=0.3,  # Adjust the height of the UDL line above the beam as needed
                      line=dict(color="green", width=5),
                      name=f"UDL ({load_value} kN/m)",
                      opacity=0.5)
        fig.add_annotation(x=udl_center, y=0.5, text=f"{load_value} kN/m", showarrow=False)
        
    fig.update_layout(title="Free Body Diagram",
                      width=1000,
                      xaxis_title="X",
                      yaxis_title="Y",
                      showlegend=True,
                      yaxis=dict(range=[-1, 1]))  # Adjust the range as needed
    st.plotly_chart(fig, use_container_width=True)

# Function to create the Word document
def create_report(report_file, moment_image_path, shear_image_path, def_figure_path):
    doc = Document()

    doc.add_heading("Beam Calculator: Force and Deflection diagram report", level=1)

    # Save Bending Moment Diagram
    doc.add_heading("Bending Moment Diagram", level=2)
    doc.add_picture(moment_image_path, width=Inches(6), height=Inches(3.5)) 
  
    # Save Shear Force Diagram
    doc.add_heading("Shear Force Diagram", level=2)
    doc.add_picture(shear_image_path, width=Inches(6), height=Inches(3.5))
   
    # Save Deflecton Diagram
    doc.add_heading("Deflection Diagram", level=2)
    doc.add_picture(def_figure_path, width=Inches(6), height=Inches(3.5))

    # Save the report file
    doc.save(report_file)

#Function to save data into excel
def save_to_excel(material_df, nodes_df, members_df, point_loads_df, udl_loads_df,results_df):
    with io.BytesIO() as buffer:
        # Create an Excel writer using xlsxwriter as the engine
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            # Write each DataFrame to a separate Excel sheet
            material_df.to_excel(writer, sheet_name='Materials', index=False)
            nodes_df.to_excel(writer, sheet_name='Nodes', index=False)
            members_df.to_excel(writer, sheet_name='Members', index=False)
            point_loads_df.to_excel(writer, sheet_name='Point_Loads', index=False)
            udl_loads_df.to_excel(writer, sheet_name='UDL_Loads', index=False)
            results_df.to_excel(writer,sheet_name='Results',index=False)
        # Get the value from the buffer
        buffer.seek(0)
        excel_data = buffer.getvalue()
    return excel_data

# Function to load input data from Excel
def load_from_excel(uploaded_file):
    # Read the Excel file into a DataFrame
    uploaded_df = pd.read_excel(uploaded_file, sheet_name=None)
    
    # Check if the required sheets exist
    required_sheets = ["Nodes", "Materials", "Members", "Point_Loads", "UDL_Loads"]
    missing_sheets = [sheet_name for sheet_name in required_sheets if sheet_name not in uploaded_df]
    
    if missing_sheets:
        st.error(f"The uploaded Excel file is missing the following sheets: {missing_sheets}")
        return None, None, None, None, None
    else:
        # Extract data from each sheet
        nodes_df = uploaded_df["Nodes"]
        material_df = uploaded_df["Materials"]
        members_df = uploaded_df["Members"]
        point_loads_df = uploaded_df["Point_Loads"]
        udl_loads_df = uploaded_df["UDL_Loads"]
        
        return nodes_df, material_df, members_df, point_loads_df, udl_loads_df

# Initialize session state for the DataFrame and graphs if they don't already exist
if 'dataframe' not in st.session_state:
    # Create a sample DataFrame
    df = pd.DataFrame()
    st.session_state['dataframe'] = df
else:
    df = st.session_state['dataframe']

if 'Free Body Diagram' not in st.session_state:
    st.session_state['Free Body Diagram'] = False

if 'Analyse & Show Results' not in st.session_state:
    st.session_state['Analyse & Show Results'] = False
    
if 'Save to Excel' not in st.session_state:
    st.session_state['Save to Excel']=False

#Define Data
# Define location and name for word report
report_file = "report/Beam_Calculation_Report.docx"

#Load image for key
image=Image.open('images/sign_convention.PNG')

#Default data for inputs
#Default Data for Nodes
node_1 = {'name': 'Node_A', 'x': 0,'y':0,'z':0,'R_Fx':True,'R_Fy':True,'R_Fz':True,'R_Mx':True,'R_My':False,'R_Mz':False}
node_2 = {'name': 'Node_B', 'x': 5,'y':0,'z':0,'R_Fx':True,'R_Fy':True,'R_Fz':True,'R_Mx':False,'R_My':False,'R_Mz':False}
node_3 = {'name': 'Node_C', 'x': 10,'y':0,'z':0,'R_Fx':True,'R_Fy':True,'R_Fz':True,'R_Mx':False,'R_My':False,'R_Mz':False}
node_4 = {'name': 'Node_D', 'x': 15,'y':0,'z':0,'R_Fx':True,'R_Fy':True,'R_Fz':True,'R_Mx':False,'R_My':False,'R_Mz':False}
nodes_df = pd.DataFrame(columns=['name','x','y','z','R_Fx','R_Fy','R_Fz','R_Mx','R_My','R_Mz'])
nodes_df.loc[len(nodes_df)] = node_1
nodes_df.loc[len(nodes_df)] = node_2
nodes_df.loc[len(nodes_df)] = node_3
nodes_df.loc[len(nodes_df)] = node_4

#Default data for Beams
member_1 = {'Name': 'M1', 'Node_LHS':'Node_A','Node_RHS':'Node_B','Material':'Material','Iy':0.01,'Iz':0.01,'J':0.1,'A':1,}
member_2 = {'Name': 'M2', 'Node_LHS':'Node_B','Node_RHS':'Node_C','Material':'Material','Iy':0.01,'Iz':0.01,'J':0.1,'A':1,}
member_3 = {'Name': 'M3', 'Node_LHS':'Node_C','Node_RHS':'Node_D','Material':'Material','Iy':0.01,'Iz':0.01,'J':0.1,'A':1,}
members_df = pd.DataFrame(columns=['Name','Node_LHS','Node_RHS','Material','Iy','Iz','J','A'])
members_df.loc[len(members_df)] = member_1
members_df.loc[len(members_df)] = member_2
members_df.loc[len(members_df)] = member_3

#Default data for Point Loads
load_1 = {'Member': 'M1', 'Load Type':'Fy','Value':-5,'x1':2.5}
load_2 = {'Member': 'M2', 'Load Type':'Fy','Value':-5,'x1':2.5}
load_3 = {'Member': 'M3', 'Load Type':'Fy','Value':-5,'x1':2.5}
point_loads_df = pd.DataFrame(columns=['Member','Load Type','Value','x1',])
point_loads_df.loc[len(point_loads_df)] = load_1
point_loads_df.loc[len(point_loads_df)] = load_2
point_loads_df.loc[len(point_loads_df)] = load_3

#Default data for UDL Loads
load_1 = {'Member': 'M1', 'Load Type':'Fy','w1':-5,'w2':-5,'x1':0,'x2':5}
load_2 = {'Member': 'M2', 'Load Type':'Fy','w1':-5,'w2':-5,'x1':0,'x2':5}
load_3 = {'Member': 'M3', 'Load Type':'Fy','w1':-5,'w2':-5,'x1':0,'x2':5}
udl_loads_df = pd.DataFrame(columns=['Member','Load Type','w1','w2','x1','x2'])
udl_loads_df.loc[len(udl_loads_df)] = load_1
udl_loads_df.loc[len(udl_loads_df)] = load_2
udl_loads_df.loc[len(udl_loads_df)] = load_3

#Default data Material Table
material_1 = {'Material':'Material','Stiffness E': 210000, 'Shear Modulus G':75000,'Poission ratio':0.3,'Density':75}
material_df = pd.DataFrame(columns=['Material','Stiffness E','Shear Modulus G','Poission ratio','Density'])
material_df.loc[len(material_df)] = material_1

#Streamlit Page configaration
#Sidebar
st.set_page_config(page_title="Beam Calculator",page_icon=":computer",layout="wide")
with st.sidebar:
    st.header('Infomation')
    st.subheader('About')
    st.write('This beam calculator has been developed for educational purposes only and any results should be independently verified.')
    st.write('The calculation are undertaken using external library Pynite. See documentation [here](https://pynite.readthedocs.io/en/latest/index.html)')
    st.subheader('Sign convention')
    st.image(image)
    st.subheader('Notes')
    st.write('When adding element to tables ensure the index added/updated. The index for each table should run from 0 and increase sequentially i.e. 0,1,2')
    st.write('Units are defined by user i.e. outputs units are respective to inputs')
    st.subheader('Planned Development')
    st.write('''
    - Add presets for Single Span or Multispan
    - Add moving loads option
    - Add Generative AI Input
    ''')

    st.subheader('Author')
    st.write('''
             - Tom Cartigny
             - Git Hub: [@TomCarts/Beam/Calculator](https://github.com/TomCarts/Beam_Calculator.git)
             ''')
    
# Main page content
st.title("Beam Calculator")
st.write('Import data from excel file or use tables below to edit inputs')

#Import Data Section
st.header("Import Data", divider='rainbow')
st.set_option('deprecation.showPyplotGlobalUse', False)

#Upload file
st.subheader('File Upload')
uploaded_file = st.file_uploader("Upload Excel file", type=["xlsx"])

if uploaded_file is not None:
    # Load data from Excel file
    nodes_df, material_df, members_df, point_loads_df, udl_loads_df = load_from_excel(uploaded_file)

#Input using dataframes section  
st.header("Tabulated Inputs", divider='rainbow')
st.subheader("Geometry", divider='grey')

col1, col2 = st.columns(2)
with col1:
    st.subheader("Nodes / Restraints")
    nodes_df=st.data_editor(nodes_df, hide_index=True, num_rows="dynamic",column_config={'y':None,'z':None})
    st.write('R_Fx/y/z - Tick to fix translation in corresponding local axis')
    st.write('R_Mx/y/z - Tick to fix rotation in corresponding local axis')
    
with col2:
    st.subheader("Members")
    members_df=st.data_editor(members_df,hide_index=True, num_rows="dynamic", column_config={'J':None})
    st.write('Data must match nodes and materials references')

st.subheader("Loads", divider='grey')
    
col1, col2 = st.columns(2)

with col1:
    st.subheader("Point Loads")
    point_loads_df=st.data_editor(point_loads_df, hide_index=True, num_rows="dynamic")
    st.write('Value - Magnitude of Load')
    st.write('Load Type: Fy - Vertical Point Loads, Mz - Moment about major axis')
    st.write('x1 - Location of point load respective to start of member i.e. from left hand node')
    
with col2:
    
    st.subheader("UDL Loads")
    udl_loads_df=st.data_editor(udl_loads_df, hide_index=True, num_rows="dynamic")
    st.write('Load Type: Fy - Vertical Point Loads, Mz - Moment about major axis')
    st.write('w1 & w2 - udl magnitude at LHS & RHS respectivly')
    st.write('x1 & x2 - udl start and finish point respective to start of member i.e. from left hand node')
    
st.subheader("Material", divider='grey')
material_df=st.data_editor(material_df, hide_index=True, num_rows="dynamic")

#Visualisation Section
st.header("Visualise", divider='rainbow')

# Button to generate the Free Body Diagram
if st.button("Free Body Diagram"):
    st.session_state['Free Body Diagram'] = True
    
if st.session_state['Free Body Diagram']: 
    plot_free_body_diagram(nodes_df, members_df, point_loads_df, udl_loads_df)

#Analysis section
st.header("Analyse", divider='rainbow')
    
# Button to analyse and display results
if st.button("Analyse & Show Results"):
    st.session_state['Analyse & Show Results'] = True

if st.session_state["Analyse & Show Results"]:
    # Perform analysis
    beam = beam_analysis(nodes_df, members_df, point_loads_df)
    
    #Compile results dataframe
    results_df = results_table(beam, members_df, nodes_df)

    #Display results section
    st.header('Results')
    
    #Display moment diagram
    st.subheader('Moment diagram')
    mom_fig = px.line(results_df, x='X', y=['beam', 'Mz'])
    moment_image_path = "images/moment_figure.png"  # Define path to save the image
    mom_fig.write_image(moment_image_path)
    st.plotly_chart(mom_fig, use_container_width=True)
    
    #Display shear force diagram
    st.subheader('Shear Force Diagram')
    shear_fig = px.line(results_df, x='X', y=['beam', 'Fy'])
    shear_image_path = "images/shear_figure.png"  # Define path to save the image
    shear_fig.write_image(shear_image_path)
    st.plotly_chart(shear_fig, use_container_width=True)
    
    # Display deflection plot
    st.subheader('Deflection Plot')
    def_fig = px.line(results_df, x='X', y=['beam', 'dy'])
    def_figure_path = "images/deflection_figure.png"
    def_fig.write_image(def_figure_path)
    st.plotly_chart(def_fig, use_container_width=True)
    
    #Display dataframe
    st.subheader('Table')
    st.dataframe(results_df)
    
    # Save to Excel button
    if st.button("Save to Excel"):
        st.session_state['Save to Excel'] = True
        
    #Download excel file button
    if st.session_state["Save to Excel"]:
        csv = save_to_excel(material_df, nodes_df, members_df, point_loads_df,udl_loads_df,results_df)
        st.download_button(
            label="Download Excel Data",
            data=csv,
            file_name='Beam_Calculator_Excel.xlsx',
            mime='text/xlsx',
            )
    # Download button for the report
    create_report(report_file, moment_image_path, shear_image_path, def_figure_path)
    st.download_button(label="Download Word Report", data=open(report_file, "rb"), file_name="Beam_Calculator_Report.docx", mime="application/octet-stream")