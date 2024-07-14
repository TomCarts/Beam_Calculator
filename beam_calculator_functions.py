from PyNite import FEModel3D
import io 
from docx.shared import Inches
import numpy as np
from docx import Document
import plotly.graph_objects as go


# Analysis function
def beam_analysis(nodes_df, members_df, point_loads_df):
# Create a new finite element model
    beam = FEModel3D()
    
    #Add material
    for i in range(len(material_df)):
        beam.add_material(material_df.iloc[i,0],material_df.iloc[i,1],material_df.iloc[i,2],material_df.iloc[i,3],material_df.iloc[i,4])
    
    # Add nodes 
    for i in range(len(nodes_df)):
        beam.add_node(nodes_df.iloc[i,0],nodes_df.iloc[i,1],nodes_df.iloc[i,2],nodes_df.iloc[i,3])

    # Add  beams:
    for i in range(len(members_df)):
        beam.add_member(members_df.iloc[i,0], members_df.iloc[i,1], members_df.iloc[i,2], members_df.iloc[i,3], members_df.iloc[i,4], members_df.iloc[i,5], members_df.iloc[i,6], members_df.iloc[i,7])

    # Provide simple supports
    for i in range(len(nodes_df)):
        beam.def_support(nodes_df.iloc[i,0], nodes_df.iloc[i,4], nodes_df.iloc[i,5], nodes_df.iloc[i,6], nodes_df.iloc[i,7], nodes_df.iloc[i,8], nodes_df.iloc[i,9])  

    # Add point loads to member
    for i in range(len(point_loads_df)):
        beam.add_member_pt_load(point_loads_df.iloc[i,0],point_loads_df.iloc[i,1],point_loads_df.iloc[i,2],point_loads_df.iloc[i,3]) # 5 kips Dead load

    #add_member_dist_load(member_name, Direction, w1, w2, x1=None, x2=None, case='Case 1')
    for i in range(len(udl_loads_df)):
        beam.add_member_dist_load(udl_loads_df.iloc[i,0],udl_loads_df.iloc[i,1],udl_loads_df.iloc[i,2],udl_loads_df.iloc[i,3],udl_loads_df.iloc[i,4],udl_loads_df.iloc[i,5]) # 5 kips Dead load4

    # Analyze the beam and perform a statics check
    beam.analyze(check_statics=True)
           
    return beam

# Display function for plots
def display_plots(beam,members_df):
    for i in range(len(members_df)):
        st.pyplot(beam.Members[members_df.iloc[i,0]].plot_shear('Fy'))
        st.pyplot(beam.Members[members_df.iloc[i,0]].plot_moment('Mz'))
        st.pyplot(beam.Members[members_df.iloc[i,0]].plot_deflection('dy'))
        
def results_table(beam,members_df,nodes_df):
    results_df=pd.DataFrame()
    Mz=[]
    Fy=[]
    dy=[]
    X=[]
    for i in range(len(members_df)):
        #for j in range((nodes_df.iloc[i+1,1]-nodes_df.iloc[i,1])*10):
        x=np.linspace(0,(nodes_df.iloc[i+1,1]-nodes_df.iloc[i,1]),100)
        for xs in x:
            Mz.append(beam.Members[members_df.iloc[i,0]].moment('Mz',xs))
            Fy.append(beam.Members[members_df.iloc[i,0]].shear('Fy',xs))
            dy.append(beam.Members[members_df.iloc[i,0]].deflection('dy',xs))
            X.append(xs+nodes_df.iloc[i,1])

    results_df['X']=X
    results_df['Mz']=Mz
    results_df['Fy']=Fy
    results_df['dy']=dy
    results_df['beam']=0  

    return results_df

# Free Body Diagram Visualization function
def plot_free_body_diagram(nodes_df, members_df, point_loads_df, udl_loads_df):
    fig = go.Figure()

    # Draw beams
    for index, row in members_df.iterrows():
        node_lhs = nodes_df[nodes_df['name'] == row['Node_LHS']]
        node_rhs = nodes_df[nodes_df['name'] == row['Node_RHS']]
        fig.add_trace(go.Scatter(x=[node_lhs['x'].values[0], node_rhs['x'].values[0]],
                                 y=[0, 0],
                                 mode='lines',
                                 line=dict(width=5),
                                 name=row['Name']))

    # Draw supports
    for index, row in nodes_df.iterrows():
        support_label = ''
        if row['R_Fx']:
            support_label += 'Fx, '
        if row['R_Fy']:
            support_label += 'Fy, '
        if row['R_Fz']:
            support_label += 'Fz, '
        if row['R_Mx']:
            support_label += 'Mx, '
        if row['R_My']:
            support_label += 'My, '
        if row['R_Mz']:
            support_label += 'Mz, '

        if support_label:
            fig.add_annotation(x=row['x'], y=-0.1, text=support_label[:-2], showarrow=False)

    # Draw point loads as arrows
    for index, load in point_loads_df.iterrows():
        member = load['Member']
        member_data = members_df[members_df['Name'] == member]
        load_x = load['x1'] + nodes_df[nodes_df['name'] == member_data['Node_LHS'].values[0]]['x'].values[0]
        load_value = load['Value']

        fig.add_annotation(x=load_x, y=0, text=f"{load_value} kN",
                           showarrow=True,
                           arrowhead=2,
                           arrowsize=1,
                           arrowwidth=2,
                           arrowcolor="blue",
                           ax=0,
                           ay=-30)

    # Draw UDLs as single thick line with label
    for index, load in udl_loads_df.iterrows():
        member = load['Member']
        member_data = members_df[members_df['Name'] == member]
        load_x1 = load['x1'] + nodes_df[nodes_df['name'] == member_data['Node_LHS'].values[0]]['x'].values[0]
        load_x2 = load['x2'] + nodes_df[nodes_df['name'] == member_data['Node_LHS'].values[0]]['x'].values[0]
        load_value = load['w1']  # Assuming w1 represents the value of UDL
        udl_center = (load_x1 + load_x2) / 2
        
        fig.add_shape(type="line",
                      x0=load_x1,
                      y0=0.3,  # Adjust the height of the UDL line above the beam as needed
                      x1=load_x2,
                      y1=0.3,  # Adjust the height of the UDL line above the beam as needed
                      line=dict(color="green", width=5),
                      name=f"UDL ({load_value} kN/m)",
                      opacity=0.5)
        fig.add_annotation(x=udl_center, y=0.5, text=f"{load_value} kN/m", showarrow=False)
        
    fig.update_layout(title="Free Body Diagram",
                      width=1000,
                      xaxis_title="X",
                      yaxis_title="Y",
                      showlegend=True,
                      yaxis=dict(range=[-1, 1]))  # Adjust the range as needed
    st.plotly_chart(fig, use_container_width=True)

# Function to create the Word document
def create_report(report_file, moment_image_path, shear_image_path, def_figure_path):
    doc = Document()

    doc.add_heading("Beam Calculator: Force and Deflection diagram report", level=1)

    # Save Bending Moment Diagram
    doc.add_heading("Bending Moment Diagram", level=2)
    doc.add_picture(moment_image_path, width=Inches(6), height=Inches(3.5)) 
  
    # Save Shear Force Diagram
    doc.add_heading("Shear Force Diagram", level=2)
    doc.add_picture(shear_image_path, width=Inches(6), height=Inches(3.5))
   
    # Save Deflecton Diagram
    doc.add_heading("Deflection Diagram", level=2)
    doc.add_picture(def_figure_path, width=Inches(6), height=Inches(3.5))

    # Save the report file
    doc.save(report_file)

#Function to save data into excel
def save_to_excel(material_df, nodes_df, members_df, point_loads_df, udl_loads_df,results_df):
    with io.BytesIO() as buffer:
        # Create an Excel writer using xlsxwriter as the engine
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            # Write each DataFrame to a separate Excel sheet
            material_df.to_excel(writer, sheet_name='Materials', index=False)
            nodes_df.to_excel(writer, sheet_name='Nodes', index=False)
            members_df.to_excel(writer, sheet_name='Members', index=False)
            point_loads_df.to_excel(writer, sheet_name='Point_Loads', index=False)
            udl_loads_df.to_excel(writer, sheet_name='UDL_Loads', index=False)
            results_df.to_excel(writer,sheet_name='Results',index=False)
        # Get the value from the buffer
        buffer.seek(0)
        excel_data = buffer.getvalue()
    return excel_data

# Function to load input data from Excel
def load_from_excel(uploaded_file):
    # Read the Excel file into a DataFrame
    uploaded_df = pd.read_excel(uploaded_file, sheet_name=None)
    
    # Check if the required sheets exist
    required_sheets = ["Nodes", "Materials", "Members", "Point_Loads", "UDL_Loads"]
    missing_sheets = [sheet_name for sheet_name in required_sheets if sheet_name not in uploaded_df]
    
    if missing_sheets:
        st.error(f"The uploaded Excel file is missing the following sheets: {missing_sheets}")
        return None, None, None, None, None
    else:
        # Extract data from each sheet
        nodes_df = uploaded_df["Nodes"]
        material_df = uploaded_df["Materials"]
        members_df = uploaded_df["Members"]
        point_loads_df = uploaded_df["Point_Loads"]
        udl_loads_df = uploaded_df["UDL_Loads"]
        
        return nodes_df, material_df, members_df, point_loads_df, udl_loads_df